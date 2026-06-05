"""Generate manager-facing Word report for gateway/simulator FIX session analysis."""
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Gateway_Simulator_FIX_Session_Analysis_2026-05-22.docx"


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
    doc.add_paragraph()


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading(
        "Technical Summary — Gateway / FINRA FIX Simulator (SP Session)", 0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.add_run("Analysis date: ").bold = True
    meta.add_run("2026-05-22\n")
    meta.add_run("Session: ").bold = True
    meta.add_run("FIX.4.4:JPMS/44B1 ↔ FNRA/SP\n")
    meta.add_run("Environment: ").bold = True
    meta.add_run(
        "Gateway (initiator) at 192.168.1.211 → Simulator (acceptor) at 192.168.1.66 "
        "ports 49100 (Primary) / 59100 (Secondary failover)"
    )
    doc.add_paragraph()

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "Connectivity and logon can succeed on the first attempt (approximately 4 seconds). "
        "The session then becomes unstable because the simulator does not send heartbeats "
        "within the gateway’s expected interval, which triggers TestRequest and "
        "disconnect/reconnect loops."
    )
    doc.add_paragraph(
        "Separately, several “Timed out waiting for logon response” events occur because "
        "the simulator’s outbound Logon is delayed beyond the gateway’s logon timeout "
        "(approximately 10 seconds) on that TCP connection—primarily due to a 60-second "
        "heartbeat delay configured in the deployed quickfixj-server.cfg, not the intended "
        "15-second logon delay."
    )

    doc.add_heading("Observed Behavior", level=1)
    add_table(
        doc,
        ["Phase", "Result"],
        [
            ("Initial logon (09:34:37 UTC)", "Success — gateway Logon and FNRA Logon within ~4 s"),
            (
                "Steady state (after ~09:35)",
                "Failure — gateway sends TestRequest, then disconnects for heartbeat timeout",
            ),
            (
                "Reconnects (09:35–09:50)",
                "Repeated gateway Logons; FNRA Logon often ~3–5 s later on a new connection",
            ),
            (
                "Example: 09:37:36 UTC",
                "Gateway Logon seq 8 — no FNRA response within 10 s; disconnect at 09:37:46; "
                "FNRA Logon ~35 s later on dead/superseded connection",
            ),
        ],
    )

    doc.add_heading("Root Causes (Technical)", level=1)

    doc.add_heading("1. Heartbeat delay blocking session processing (primary)", level=2)
    for bullet in [
        "Deployed config runs HeartBtDelay=Y with HeartBtDelayTime=60 (60 seconds), not 3 seconds in the source template.",
        "While waiting in toAdmin, the simulator does not send heartbeats and ignores inbound admin messages (gateway heartbeats and TestRequest).",
        "Gateway expects traffic within ~1.5 × HeartBtInt (30 s) ≈ 45 s; after silence it sends TestRequest, then disconnects (“Timed out waiting for heartbeat”).",
        "The same delay blocks outbound Logon on the session thread, so reconnect logons miss the ~10 s logon timeout.",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("2. Configuration mismatch (deployed vs intended)", level=2)
    for bullet in [
        "Simulator loads ./quickfixj-server.cfg from the server working directory (overrides JAR).",
        "Logs show for SP: effective LogonDelay=false (LogonDelay=N on [session]); LogonDelayinSecs=15 is not active.",
        "Startup shows HeartBtDelay=false at [default] but 60 s delay at runtime on the session block.",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("3. Dual-instance / shared database", level=2)
    for bullet in [
        "Primary and Secondary simulators run together against the same TRACE_FIX_SESSIONS JDBC store.",
        "Gateway alternates ports 49100 / 59100 (failover). Both instances update shared sequence numbers.",
        "Results: gaps, ResendRequest, SequenceReset, and slower logon handling.",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("4. Per-connection vs end-to-end log interpretation", level=2)
    for bullet in [
        "Message logs can show FNRA Logon eventually after many gateway Logons (looks like all are answered).",
        "Event-log timeouts apply per TCP connection: if FNRA replies after the gateway disconnected and reconnected, that still counts as logon timeout on the failed attempt.",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("Evidence (Key Timestamps, UTC)", level=1)
    add_table(
        doc,
        ["Event", "Time", "Detail"],
        [
            ("Successful logon", "09:34:37 → 09:34:41", "Normal acceptor response"),
            ("Heartbeat stall", "09:35:08 HB → 09:35:23 TEST", "~46 s after logon; no timely FNRA HB"),
            ("Logon timeout (example)", "09:37:36 → 09:37:46", "No FNRA Logon within 10 s"),
            ("Late FNRA Logon", "09:38:11", "After 60 s HeartBtDelay; too late for prior connect"),
        ],
    )
    doc.add_paragraph(
        "Simulator log (aligned): “HeartBtDelay=Y: waiting 60s before sending Heartbeat” and "
        "“Ignoring admin message until we send Heartbeat back.”"
    )

    doc.add_heading("Example: 09:37:46 Logon Timeout", level=1)
    add_table(
        doc,
        ["Step", "Time (UTC)", "Detail"],
        [
            ("Gateway sends Logon", "09:37:36.325", "MsgSeqNum 8"),
            ("Gateway disconnects", "09:37:46.333", "Timed out waiting for logon response (~10 s)"),
            ("Simulator receives Logon", "~09:37:39", "fromAdmin processes inbound"),
            ("Simulator sends FNRA Logon", "09:38:11.797", "~35 s late; gateway already disconnected"),
        ],
    )
    doc.add_paragraph(
        "The waiting logon was gateway Logon 34=8 on that TCP connect. Later FNRA Logon (e.g. 34=9 at "
        "09:38:11) answers a new reconnect, not the failed 34=8 attempt."
    )

    doc.add_heading("Gateway Logon vs FNRA Response (Summary)", level=1)
    add_table(
        doc,
        ["Gateway Logon (UTC)", "Seq", "FNRA response within ~15 s?", "Notes"],
        [
            ("09:34:37.343", "1", "Yes → 09:34:41", "OK"),
            ("09:35:51.327", "4", "Yes → 09:35:54", "OK (~3.6 s)"),
            ("09:37:05.332", "7", "No", "Likely timeout; new Logon follows"),
            ("09:37:36.325", "8", "No (before 09:37:46)", "09:37:46 timeout event"),
            ("09:38:06.333", "9", "Yes → 09:38:11", "New connect after 34=8 failed"),
        ],
    )

    doc.add_heading("Impact", level=1)
    for bullet in [
        "Gateway cannot maintain a stable logged-on session for SP.",
        "TRACE FIX testing over this session is blocked or flaky.",
        "Failover testing is confused by two simulators and sequence resets.",
        "Timeouts may be misattributed to “logon delay 15 s” while the active issue is heartbeat delay 60 s and logon timeout ~10 s.",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("Recommended Actions", level=1)
    add_table(
        doc,
        ["Priority", "Action"],
        [
            (
                "P1",
                "On server quickfixj-server.cfg: set HeartBtDelay=N for production keepalive, "
                "or HeartBtDelayTime=3 (not 60). Restart one simulator instance.",
            ),
            ("P1", "Confirm SP session: HeartBeat_Required=Y."),
            (
                "P2",
                "If testing logon delay: LogonDelay=Y, LogonDelayinSecs=15 on SP [session]; "
                "gateway LogonTimeout ≥ 25–30 s.",
            ),
            (
                "P2",
                "For single-environment tests: run Primary only or Secondary only, not both on same DB.",
            ),
            (
                "P3",
                "Re-test: stable session >5 min, no TEST/disconnect loop; logon retry within gateway timeout.",
            ),
        ],
    )

    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph(
        "The simulator application and ThreadedSocketAcceptor behave as designed; failures are driven by "
        "deployed configuration (60 s heartbeat delay holding the session thread) and operational setup "
        "(dual instance + shared sequences). Fixing heartbeat delay and aligning config with test intent "
        "should resolve both heartbeat timeouts and most logon response timeouts. Logon delay testing "
        "requires explicit LogonDelay=Y on the session and a gateway logon timeout longer than the "
        "configured delay."
    )

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run("Prepared from correlated gateway message log and simulator.log (2026-05-22).").italic = True
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
