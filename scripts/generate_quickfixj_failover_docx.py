"""Generate Word guide: QuickFIX/J session failover and admin-message behavior."""
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "QuickFIXJ_Session_Failover_and_Admin_Messages.docx"


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()


def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet" if level == 0 else "List Bullet 2")
    return p


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading(
        "QuickFIX/J Session Failover & Admin Messages — Reference Guide", 0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.add_run("Document date: ").bold = True
    meta.add_run(f"{date.today().isoformat()}\n")
    meta.add_run("Scope: ").bold = True
    meta.add_run(
        "QuickFIX/J engine behavior (acceptor/initiator failover, sequence recovery) "
        "plus FINRA TRACE Simulator Primary/Secondary deployment.\n"
    )
    meta.add_run("Sources: ").bold = True
    meta.add_run(
        "QuickFIX/J Overview & Configuration (quickfixj.org/docs), "
        "Acceptor Failover user manual, project configs and docs/RESENDREQUEST_AND_SEQUENCE_GAPS.md."
    )
    doc.add_paragraph()

    # --- 1 Overview ---
    doc.add_heading("1. What QuickFIX/J Does for You", level=1)
    doc.add_paragraph(
        "QuickFIX/J is a Java FIX engine that owns session lifecycle: TCP connect/disconnect, "
        "Logon/Logout negotiation, MsgSeqNum (tag 34) tracking, heartbeat scheduling, gap detection, "
        "ResendRequest generation, and persistence via a MessageStore (file or JDBC). "
        "Your application implements Application callbacks (fromAdmin, toAdmin, fromApp, toApp); "
        "the engine handles protocol-level admin messages unless you reject or modify them."
    )
    doc.add_paragraph(
        "Official entry point: QuickFIX/J Overview — "
        "https://quickfixj.org/docs/overview"
    )

    # --- 2 Session identity ---
    doc.add_heading("2. Session Identity (What Defines “One Session”)", level=1)
    doc.add_paragraph(
        "A FIX session is identified by BeginString + Sender/Target CompID (and optional SubID, "
        "LocationID, SessionQualifier on initiators). The TCP port is NOT part of the session key. "
        "Therefore Primary acceptor (port 64034) and Secondary acceptor (port 64134) with the same "
        "SenderCompID/TargetCompID share one logical session if they use the same persistent store."
    )
    add_table(
        doc,
        ["Component", "Acceptor (simulator)", "Initiator (gateway)"],
        [
            ("BeginString", "FIX.4.4", "FIX.4.4"),
            ("SenderCompID", "FNRA (+ SubID SP/CA/TS)", "JPMS (+ SubID 44B1/44B2/44B3)"),
            ("TargetCompID", "JPMS", "FNRA"),
            ("Store key", "TRACE_FIX_SESSIONS row", "Same row when JDBC shared"),
        ],
    )

    # --- 3 Failover models ---
    doc.add_heading("3. Failover Models in QuickFIX/J", level=1)

    doc.add_heading("3.1 Acceptor-side failover (passive standby)", level=2)
    doc.add_paragraph(
        "Two (or more) acceptor JVMs listen on different ports but share a persistent MessageStore "
        "(FileStore on shared disk, or JdbcStore on shared DB). Only one acceptor holds the live "
        "TCP session at a time. When the active node dies, the initiator reconnects to the standby "
        "port. On the first Logon to the standby, QuickFIX/J must reload sequence state from the store."
    )
    bullet(doc, "Required setting: RefreshOnLogon=Y (default N).")
    bullet(
        doc,
        "Documented in QuickFIX/J “Simple Failover for Socket Acceptors” and "
        "“Configuring Session Failover” wiki.",
    )
    bullet(
        doc,
        "This project: Primary (quickfixj-server.cfg) + Secondary (quickfixj-server-secondary.cfg), "
        "shared TRACE_FIX_SESSIONS / TRACE_FIX_MESSAGES, RefreshOnLogon=Y, UseJdbcStore=Y.",
    )

    doc.add_heading("3.2 Initiator-side failover (client dials backup)", level=2)
    doc.add_paragraph(
        "On the initiator, QuickFIX/J supports multiple connect targets. If the primary host/port "
        "fails, the engine tries backup addresses on ReconnectInterval until one succeeds."
    )
    add_table(
        doc,
        ["Setting", "Purpose"],
        [
            ("SocketConnectHost / SocketConnectPort", "Primary target"),
            ("SocketConnectHost1 / SocketConnectPort1", "First backup (example from QFJ sample cfg)"),
            ("SocketConnectHost2 / SocketConnectPort2", "Second backup"),
            ("ReconnectInterval", "Seconds between reconnect attempts (initiator only)"),
            ("LogonTimeout", "Seconds to wait for Logon response before disconnect (default 10)"),
            ("LogoutTimeout", "Seconds to wait for Logout response (default 2)"),
        ],
    )
    doc.add_paragraph(
        "Gateway must be configured to try Simulator Primary ports first, then Secondary ports "
        "(e.g. 64034 → 64134 for SP). Exact tag names depend on the gateway’s FIX engine."
    )

    # --- 4 Timeline ---
    doc.add_heading("4. End-to-End Timeline When Primary Fails", level=1)
    phases = [
        (
            "T0 — Steady state on Primary",
            "TCP connected; both sides logged on; HeartBtInt agreed in Logon; engine sends/receives "
            "Heartbeat (35=0) on schedule; application messages use incrementing MsgSeqNum; "
            "TRACE_FIX_SESSIONS updated on each persisted message.",
        ),
        (
            "T1 — Primary process or host dies",
            "TCP breaks (RST or timeout). Initiator detects disconnect. Acceptor on Primary is gone — "
            "no Logout may be sent. With ResetOnDisconnect=N (simulator default), sequences are NOT "
            "reset in the store.",
        ),
        (
            "T2 — Initiator reconnect loop",
            "Initiator waits ReconnectInterval, retries Primary; if unreachable, tries Secondary host/port. "
            "May open new TCP to Secondary while still holding last known seq nums in its store.",
        ),
        (
            "T3 — Logon on Secondary (critical)",
            "Initiator sends Logon (35=A) with its next MsgSeqNum (e.g. 150). Secondary acceptor: "
            "(1) RefreshOnLogon loads incoming_seqnum/outgoing_seqnum from JDBC; "
            "(2) compares received 34 vs expected; "
            "(3) if gap: sends ResendRequest (35=2) before accepting new app traffic; "
            "(4) responds with Logon (EncryptMethod, HeartBtInt, optional 789); session is logged on.",
        ),
        (
            "T4 — Post-failover steady state",
            "Heartbeats resume; seq nums continue from shared DB (no reset unless 141=Y or admin "
            "SequenceReset). TRACE_FIX_MESSAGES_LOG shows simulator_instance = FixSimulator-Secondary.",
        ),
        (
            "T5 — Primary returns (failback)",
            "Only one TCP should be active. If gateway failbacks to Primary while Secondary still had "
            "the session, same RefreshOnLogon rules apply. Running BOTH JVMs connected to the same "
            "gateway session is an operational error (split-brain, duplicate Logons, seq fights).",
        ),
    ]
    for name, detail in phases:
        p = doc.add_paragraph()
        p.add_run(name + ": ").bold = True
        p.add_run(detail)

    # --- 5 When to expect Logon/Logout ---
    doc.add_heading("5. When to Expect Logon and Logout", level=1)

    doc.add_heading("5.1 Logon (35=A)", level=2)
    add_table(
        doc,
        ["Situation", "Who sends Logon first", "What you should see"],
        [
            ("Cold start / first connect of the day", "Typically initiator", "Both sides seq often 1 unless store has state"),
            ("Reconnect after TCP drop", "Typically initiator", "Logon with next seq from initiator store; acceptor RefreshOnLogon"),
            ("Acceptor failover (Primary → Secondary)", "Initiator to new port", "One Logon pair; acceptor loads DB; may include 789"),
            ("Scheduled session (StartTime/EndTime)", "After boundary reset if configured", "NonStopSession=Y avoids daily window reject"),
            ("Simulator LogonDelay=Y", "Initiator first; acceptor delays outbound Logon", "Risk of LogonTimeout if delay > gateway timeout"),
        ],
    )

    doc.add_heading("5.2 Logout (35=5)", level=2)
    add_table(
        doc,
        ["Situation", "Typical behavior", "Seq impact (simulator ResetOnLogout=N)"],
        [
            ("Graceful shutdown (either side)", "Logout then disconnect", "Seq preserved unless ResetOnLogout=Y"),
            ("Engine disconnect due to seq error", "Logout with Text= reason, then disconnect", "No auto reset (N)"),
            ("Heartbeat timeout", "Disconnect; Logout may or may not be exchanged", "No reset on disconnect (N)"),
            ("Simulator SendLogout_at_Shutdown=Y", "Acceptor sends Logout to all sessions on JVM exit", "App persists outgoing_seqnum+1 to DB"),
            ("Initiator Logout received", "Acceptor processes; simulator may persist next sender", "Next Logon should use N+1"),
        ],
    )

    # --- 6 Admin messages ---
    doc.add_heading("6. Admin Messages (MsgType / Tag 35)", level=1)
    doc.add_paragraph(
        "Admin messages are handled by the QuickFIX/J Session before (or via) Application callbacks. "
        "During failover, these are the messages you will see in .event.log and .messages.log."
    )
    add_table(
        doc,
        ["MsgType", "Name", "Direction / trigger", "Failover relevance"],
        [
            ("0", "Heartbeat", "Either side; every HeartBtInt", "Proves session alive after failover"),
            ("1", "TestRequest", "Side that suspects dead peer (TestReqID 112)", "Common during reconnect stress"),
            ("2", "ResendRequest", "Gap detected (BeginSeqNo 7, EndSeqNo 16)", "After failover if seq mismatch"),
            ("3", "Reject", "Session-level reject of bad admin msg", "Bad Logon, garbled msg, etc."),
            ("4", "SequenceReset", "Gap fill or full reset (NewSeqNo 36, GapFillFlag 123)", "Initiator may jump seq; updates store"),
            ("5", "Logout", "Graceful or forced end", "May precede reconnect to other port"),
            ("A", "Logon", "Session establishment", "RefreshOnLogon runs here on acceptor"),
            ("j", "BusinessMessageReject", "App-level (sometimes logged as admin path)", "After logon, bad app message"),
        ],
    )

    doc.add_heading("6.1 Heartbeat (0) and TestRequest (1)", level=2)
    doc.add_paragraph(
        "After Logon, each side sends Heartbeat if no other traffic within HeartBtInt seconds. "
        "If a Heartbeat is overdue, QuickFIX/J sends TestRequest; if still no reply, disconnects "
        "(unless DisableHeartBeatCheck=Y). HeartBeatTimeoutMultiplier and TestRequestDelayMultiplier "
        "scale the timeouts (see Configuration doc)."
    )
    bullet(doc, "Expect Heartbeats within ~HeartBtInt after successful failover Logon.")
    bullet(doc, "Missing HB → TestRequest → disconnect is independent of failover path.")

    doc.add_heading("6.2 ResendRequest (2)", level=2)
    doc.add_paragraph(
        "Sent when ValidateSequenceNumbers=Y and incoming MsgSeqNum > next expected. "
        "BeginSeqNo/EndSeqNo define the gap. Engine sends ONE request per gap by default; "
        "SendRedundantResendRequests=Y allows repeat requests if the gap is not filled."
    )
    bullet(doc, "After acceptor restart without RefreshOnLogon: spurious ResendRequest(1, N-1).")
    bullet(doc, "After failover with shared JDBC + RefreshOnLogon=Y: usually NO spurious resend if DB correct.")
    bullet(doc, "Counterparty must resend with PossDupFlag=Y (43) and OrigSendingTime (122).")
    bullet(doc, "Log: “Already sent ResendRequest FROM: x TO: y. Not sending another.” = gap still open.")

    doc.add_heading("6.3 SequenceReset (4)", level=2)
    doc.add_paragraph(
        "GapFillFlag=Y: advance expected incoming to NewSeqNo without resetting outgoing. "
        "GapFillFlag=N (or absent): full reset — both directions align to NewSeqNo semantics per FIX rules."
    )
    bullet(doc, "Often sent in response to ResendRequest to skip replay of admin-only gaps.")
    bullet(doc, "Engine persists new seq to MessageStore (JdbcStore → TRACE_FIX_SESSIONS).")

    doc.add_heading("6.4 Logon (A) — key tags during failover", level=2)
    add_table(
        doc,
        ["Tag", "Name", "Role in failover"],
        [
            ("34", "MsgSeqNum", "Must align with peer expectation; drives ResendRequest"),
            ("98", "EncryptMethod", "0 = none (typical sim)"),
            ("108", "HeartBtInt", "Negotiated HB interval"),
            ("141", "ResetSeqNumFlag", "Y = both sides reset to 1 (honored even if ResetOnLogon=N)"),
            ("789", "NextExpectedMsgSeqNum", "If EnableNextExpectedMsgSeqNum=Y, sync next expected on peer"),
        ],
    )
    doc.add_paragraph(
        "Simulator (WizFixApplication): on inbound Logon, loads TRACE_FIX_SESSIONS, merges DB seq with "
        "engine/file store (Secondary file store may start at 1), adjusts setNextTargetMsgSeqNum for "
        "correct 789 in outbound Logon, and logs shared-store / failover lines."
    )

    doc.add_heading("6.5 Reject (3) and Logout (5)", level=2)
    doc.add_paragraph(
        "Reject: session-level; Text (58) explains failure (e.g. MsgSeqNum too low/high, comp ID mismatch). "
        "Logout: often follows Reject or timeout; may include Text. ResetOnLogout=N keeps seq in store for reconnect."
    )

    # --- 7 Sequence settings ---
    doc.add_heading("7. Sequence Number Settings (Configuration)", level=1)
    add_table(
        doc,
        ["Setting", "Default", "Simulator", "Failover note"],
        [
            ("RefreshOnLogon", "N", "Y", "Mandatory for JDBC failover — reload DB on Logon"),
            ("ResetOnLogon", "N", "N", "Do not auto-reset; still honor peer 141=Y"),
            ("ResetOnLogout", "N", "N", "Seq survives graceful logout"),
            ("ResetOnDisconnect", "N", "N", "Seq survives TCP drop — required for FINRA-style failover"),
            ("ResetOnError", "N", "N", "No full wipe on engine error"),
            ("PersistMessages", "Y", "Y", "Required for resend replay from store"),
            ("UseJdbcStore", "—", "Y", "Shared TRACE_FIX_* tables Primary/Secondary"),
            ("SendRedundantResendRequests", "N", "Y", "Retry ResendRequest if gap stuck"),
            ("EnableNextExpectedMsgSeqNum", "N", "Y", "789 on Logon for sync"),
            ("ValidateSequenceNumbers", "Y", "Y", "Enables gap detection and ResendRequest"),
        ],
    )

    # --- 8 Decision tree ---
    doc.add_heading("8. Decision Tree: What Happens on Logon After Failover?", level=1)
    doc.add_paragraph(
        "Use this when reading event logs on the standby acceptor."
    )
    steps = [
        "Initiator connects to Secondary port → TCP established.",
        "Initiator sends Logon(MsgSeqNum = I).",
        "Acceptor RefreshOnLogon → MessageStore.refresh() → load incoming_seqnum = E, outgoing_seqnum = O from DB.",
        "If I > E → gap → ResendRequest(E, I-1) [engine]; initiator must resend or send SequenceReset.",
        "If I < E → “MsgSeqNum too low” → Logout/disconnect (unless 789/sync fixes expectation).",
        "If I == E (or after gap fill) → acceptor sends Logon(MsgSeqNum = O or O') with HeartBtInt, optional 789 = f(E,I).",
        "Session logged on → Heartbeats → application traffic continues.",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}", style="List Number")

    # --- 9 Initiator vs acceptor ---
    doc.add_heading("9. Initiator vs Acceptor Responsibilities", level=1)
    add_table(
        doc,
        ["Event", "Initiator (gateway)", "Acceptor (simulator)"],
        [
            ("TCP down", "Detects; schedules reconnect; may try backup port", "Process gone — no action"),
            ("Reconnect", "Opens TCP; sends Logon", "Accepts TCP; RefreshOnLogon; responds Logon"),
            ("Seq gap (peer ahead)", "Receives ResendRequest; resends from store", "Sends ResendRequest"),
            ("Seq gap (peer behind)", "May receive Logout/Reject", "Rejects low seq"),
            ("141=Y on Logon", "Requests reset", "Resets store to 1 if honored"),
            ("Heartbeat miss", "TestRequest / disconnect", "Same"),
        ],
    )

    # --- 10 Project deployment ---
    doc.add_heading("10. FINRA Simulator Deployment Checklist", level=1)
    bullets = [
        "Primary and Secondary: same JdbcURL and table names; only ports and FileStorePath differ.",
        "Do not point Secondary at a different database — sequences will diverge on failover.",
        "Run only one active TCP per session ID; standby listens idle until failover.",
        "Gateway: Primary host/ports (64034, 64093, 64094) then Secondary (64134, 64193, 64194).",
        "Verify startup log: SHARED SEQ DB line when UseJdbcStore=Y.",
        "Avoid running both JVMs against the same gateway simultaneously (causes ResendRequest/SequenceReset storms).",
        "After simulator restart: RefreshOnLogon=Y prevents “expecting 1” when DB has higher seq.",
        "LogonDelay / HeartBtDelay simulator features: ensure gateway LogonTimeout and HB tolerance exceed delays.",
    ]
    for b in bullets:
        bullet(doc, b)

    # --- 11 Micro-details / log lines ---
    doc.add_heading("11. Event Log Phrases (Micro-Details)", level=1)
    add_table(
        doc,
        ["Log line (typical)", "Meaning"],
        [
            ("Refreshing message/state store at logon", "RefreshOnLogon=Y in effect"),
            ("Received logon / Responding to Logon request", "Normal admin path"),
            ("Sent ResendRequest FROM: a TO: b", "Gap detected; waiting for resend"),
            ("Already sent ResendRequest... Not sending another", "SendRedundantResendRequests=N or same gap"),
            ("MsgSeqNum too high, expecting X but received Y", "Gap not filled; peer ahead"),
            ("MsgSeqNum too low, expecting X but received Y", "Peer behind; often 789 mismatch"),
            ("Timed out waiting for logon response", "Initiator LogonTimeout — peer Logon not received in time"),
            ("Timed out waiting for heartbeat", "HB/TestRequest path failed"),
            ("Disconnecting: Timed out waiting for logon response", "Initiator-side disconnect reason"),
            ("Logon attempt not within session time", "StartTime/EndTime — use NonStopSession=Y for 24/7"),
            ("Falling back to file store", "JDBC misconfigured — failover seq will be wrong"),
        ],
    )

    # --- 12 References ---
    doc.add_heading("12. References", level=1)
    refs = [
        "QuickFIX/J Overview — https://quickfixj.org/docs/overview",
        "QuickFIX/J Configuration — https://quickfixj.org/docs/configuration",
        "Configuring Session Failover (wiki) — https://www.quickfixj.org/wiki/display/qfj/Configuring+Session+Failover",
        "Acceptor Failover (user manual 2.3.0) — https://www.quickfixj.org/usermanual/2.3.0/usage/acceptor_failover.html",
        "Session Javadoc (RefreshOnLogon, reset settings) — https://www.quickfixj.org/javadoc/2.3.0/quickfix/Session.html",
        "Project: PRIMARY-SECONDARY.md, GATEWAY-CONNECT-TO-SIMULATOR.md, docs/RESENDREQUEST_AND_SEQUENCE_GAPS.md",
        "Project configs: quickfixj-server.cfg, quickfixj-server-secondary.cfg",
    ]
    for r in refs:
        bullet(doc, r)

    doc.add_paragraph()
    doc.add_paragraph(
        "This document is intended for operators and developers integrating gateway initiators "
        "with the Wizcom FINRA TRACE FIX Simulator Primary/Secondary pair."
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
